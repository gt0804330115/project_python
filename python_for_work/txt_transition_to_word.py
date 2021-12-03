#!/usr/bin/env python
# -*- coding: UTF-8 -*-
"""
@Project   ：project_python
@File      ：txt_transition_to_word.py
@Author    ：guotong
@Date      ：2021/10/8 9:53
@Features  ： WORD格式转换
    1、文件名称改成标题名
    2、标题格式： 宋体 二号 加粗 居中对齐
    3、一级标题： 黑体 三号       例：一、二、三、
    4、二级标题： 楷体_GB2312 三号        例：（一）（二）（三）
    5、正文： 仿宋_GB2312 三号 首行缩进2字符 行距固定值30磅
    6、末尾一行： 右对齐

"""

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from docx import Document
import re

doc = Document('H:\工作汇总\测试.docx')
all_paragraphs = len(doc.paragraphs)
#  正则表达式
pattern01 = r'(一、)|(二、)|(三、)|(四、)|(五、)'
pattern02 = r'(（一）)|(（二）)|(（三）)|(（四）)|(（五）)'
# doc.styles['Normal'].font.name = '楷体_GB2312'
for paragraph in doc.paragraphs:  # 遍历段落
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置段落的对齐方式（两端对齐）
    paragraph.paragraph_format.first_line_indent = 406400  # 设置段落的首行缩进2字符
    paragraph.paragraph_format.line_spacing = Pt(30)  # 设置行距固定值30磅
    for run in paragraph.runs:
        match01 = re.search(pattern01, run.text)
        match02 = re.search(pattern02, run.text)
        if match01 != None:
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(16)
        elif match02 != None:
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
            run.font.size = Pt(16)
        else:
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            run.font.size = Pt(16)
doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in doc.paragraphs[0].runs:
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(22)
doc.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in doc.paragraphs[1].runs:
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    run.font.size = Pt(16)
i = 1
while True:
    if len(doc.paragraphs[all_paragraphs - i].text) != 0:
        doc.paragraphs[all_paragraphs - i].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        break
    else:
        p = doc.paragraphs[all_paragraphs - i]._element  # 删除空白行
        p.getparent().remove(p)
        doc.paragraphs[all_paragraphs - i - 1]._p = doc.paragraphs[all_paragraphs - i - 1]._element = None
        i += 1

doc.save('H:\工作汇总\测试1.docx')
