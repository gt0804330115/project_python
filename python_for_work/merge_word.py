#!/usr/bin/env python
# -*- coding: UTF-8 -*-
"""
@Project   ：project_python 
@File      ：merge_word.py
@Author    ：guotong
@Date      ：2021/10/22 9:00 
@Features  ： 合并word文档 （doc转docx）
"""
import fnmatch
import os
from docx import Document
from docxcompose.composer import Composer
import send2trash
from win32com import client as wc


# 合并docx文档
def merge_doc(source_file_path_list, target_file_path):
    '''
    合并多个docx文件
    :param source_file_path_list: 源文件路径列表
    :param target_file_path: 目标文件路径
    :return:
    '''
    # 填充分页符号文档
    page_break_doc = Document()
    page_break_doc.add_page_break()
    # 定义新文档
    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(source_file_path_list)):
        # 跳过第一个作为模板的文件
        if i == 0:
            continue
        # 填充分页符文档
        target_composer.append(page_break_doc)
        # 拼接文档内容
        f = source_file_path_list[i]
        target_composer.append(Document(f))
    # 保存目标文档
    target_composer.save(target_file_path)


# 需要检查一个“目录”，或者某个包含子目录的目录树，并根据某种模式迭代所有的文件（也可能包含子目录）使用的是os.walk生成器
def all_files(root, patterns='*', single_level=False, yield_folders=False):
    patterns = patterns.split(';')
    for path, subdirs, files in os.walk(root):
        if yield_folders:
            files.extend(subdirs)
        files.sort()
        for name in files:
            for pattern in patterns:
                if fnmatch.fnmatch(name, pattern):
                    yield os.path.join(path, name)
                    break
        if single_level:
            break


def doc_to_docx(file):
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(file)
    doc.SaveAs("{}x".format(file), 16)
    doc.Close()
    return file


if __name__ == '__main__':
    for path in all_files('E:\\新建文件夹1111', '*.doc'):
        print(path)
        if '~$' in path:
            send2trash.send2trash(path)  # 删除临时文件（回收站）
        else:
            doc_to_docx(path)
    files = []
    for path in all_files('E:\\新建文件夹1111', '*.docx'):
        files.append(path)
        merge_doc(files, 'H:\\工作汇总\\测试1.docx')
    print('已完成')
